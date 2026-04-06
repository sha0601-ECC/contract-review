import io
import base64
from typing import Optional

import pdfplumber
from docx import Document
from PIL import Image


async def parse_contract(
    content: bytes,
    suffix: str,
    filename: str,
) -> dict:
    """Parse contract file and extract text + images."""
    if suffix == ".pdf":
        return await _parse_pdf(content)
    elif suffix in [".docx", ".doc"]:
        return await _parse_docx(content, filename)
    elif suffix == ".txt":
        text = content.decode("utf-8", errors="replace")
        return {"text": text, "images": [], "filename": filename}
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


async def _parse_pdf(content: bytes) -> dict:
    """Extract text and images from PDF."""
    text_parts = []
    images = []

    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page_num, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[第{page_num + 1}页]\n{page_text}")

            # Extract images from page
            for img_info in page.images:
                try:
                    # Re-extract image from page
                    x0 = img_info.get("x0", 0)
                    y0 = img_info.get("top", 0)
                    x1 = img_info.get("x1", 0)
                    y1 = img_info.get("bottom", 0)

                    # Get image data
                    img_data = img_info.get("stream", {})
                    if hasattr(img_data, "get_data"):
                        img_bytes = img_data.get_data()
                        img_b64 = base64.b64encode(img_bytes).decode()
                        images.append(img_b64)
                except Exception:
                    # Skip images we can't extract
                    pass

    full_text = "\n\n".join(text_parts)
    return {"text": full_text, "images": images, "filename": ""}


async def _parse_docx(content: bytes, filename: str) -> dict:
    """Extract text and images from Word document."""
    doc = Document(io.BytesIO(content))
    text_parts = []

    # Extract images
    images = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_bytes = rel.target_part.blob
                img_b64 = base64.b64encode(img_bytes).decode()
                images.append(img_b64)
            except Exception:
                pass

    # Extract text
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract table text
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n\n".join(text_parts)
    return {"text": full_text, "images": images, "filename": filename}
