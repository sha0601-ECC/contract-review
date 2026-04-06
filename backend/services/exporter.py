import base64
import re
import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def html_to_docx(html_content: str, images: list[str]) -> bytes:
    """
    Convert HTML content + images to a proper .docx file.
    Images are expected as base64 encoded strings.
    """
    doc = Document()

    # Set default font for Chinese
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(
        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}eastAsia',
        'SimSun'
    )

    # Image map: placeholder -> base64 data
    # In HTML, images are embedded as <img src="data:image/png;base64,..." />
    image_map: dict[str, str] = {}
    img_counter = [0]

    def replace_img(match):
        full_src = match.group(1)
        img_counter[0] += 1
        placeholder = f'__IMG_{img_counter[0]}__'
        image_map[placeholder] = full_src
        return placeholder

    # Find all base64 images in HTML
    processed_html = re.sub(r'<img[^>]+src="(data:image/[^"]+)"', replace_img, html_content)

    # Strip HTML tags but preserve paragraph structure
    lines = processed_html.split('\n')
    current_paragraph = doc.add_paragraph()
    current_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for line in lines:
        stripped = line.strip()
        if not stripped:
            # Empty line = new paragraph
            current_paragraph = doc.add_paragraph()
            continue

        # Check if line starts a heading
        if stripped.startswith('<h1') or stripped.startswith('<h2') or stripped.startswith('<h3'):
            # Extract heading level and text
            h_match = re.match(r'<h([1-3])[^>]*>(.*)</h[1-3]>', stripped, re.DOTALL)
            if h_match:
                level = int(h_match.group(1))
                text = re.sub(r'<[^>]+>', '', h_match.group(2))
                heading = doc.add_heading(text, level=level)
                current_paragraph = heading
        elif stripped.startswith('<li') or stripped.startswith('<ul') or stripped.startswith('<ol'):
            # List items
            li_text = re.sub(r'<[^>]+>', '', stripped)
            if li_text:
                para = doc.add_paragraph(li_text, style='List Bullet')
                current_paragraph = para
        elif stripped.startswith('<blockquote'):
            # Blockquote
            bq_text = re.sub(r'<[^>]+>', '', stripped)
            if bq_text:
                para = doc.add_paragraph(bq_text)
                para.paragraph_format.left_indent = Inches(0.3)
                para.font.italic = True
                current_paragraph = para
        else:
            # Regular paragraph - remove remaining HTML tags
            text = re.sub(r'<[^>]+>', '', stripped)
            text = text.replace('&nbsp;', ' ')
            text = text.replace('&lt;', '<')
            text = text.replace('&gt;', '>')
            text = text.replace('&amp;', '&')

            # Check for image placeholders inline
            while '__IMG_' in text:
                start = text.find('__IMG_')
                end = text.find('__', start + 1)
                if end == -1:
                    break
                placeholder = text[start:end]
                img_b64 = image_map.get(placeholder, '')
                text = text[:start] + text[end + 2:]

                if img_b64 and img_b64.startswith('data:image'):
                    # Extract base64 data
                    comma_idx = img_b64.find(',')
                    if comma_idx != -1:
                        b64_data = img_b64[comma_idx + 1:]
                        try:
                            img_bytes = base64.b64decode(b64_data)
                            # Add image to document
                            img_stream = io.BytesIO(img_bytes)
                            current_paragraph.add_run()
                            run = current_paragraph.runs[-1]
                            run.add_picture(img_stream, width=Inches(4))
                        except Exception:
                            pass

            if text.strip():
                if current_paragraph.text:
                    current_paragraph.add_run(' ' + text)
                else:
                    current_paragraph.add_run(text)

    # Handle block-level images (images on their own line)
    for placeholder, img_b64 in image_map.items():
        if not img_b64.startswith('data:image'):
            continue

        comma_idx = img_b64.find(',')
        if comma_idx == -1:
            continue
        b64_data = img_b64[comma_idx + 1:]

        try:
            img_bytes = base64.b64decode(b64_data)
            img_stream = io.BytesIO(img_bytes)
            para = doc.add_paragraph()
            run = para.add_run()
            run.add_picture(img_stream, width=Inches(5))
        except Exception:
            pass

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
